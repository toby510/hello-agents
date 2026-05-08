"""
Word 文档 → 文本分块 → Qdrant 向量数据库
完整 RAG 核心流程，可一次性执行

依赖安装：
    pip install python-docx qdrant-client sentence-transformers tqdm

运行方式：
    python docx_to_qdrant.py
    python docx_to_qdrant.py --file my_doc.docx --collection my_col --chunk-size 500
"""

import argparse
import hashlib
import re
import sys
import uuid
from dataclasses import dataclass, field
from pathlib import Path
from typing import Generator

# ── 依赖检查 ──────────────────────────────────────────────────────────────────
def _check_dependencies():
    missing = []
    for pkg, import_name in [
        ("python-docx",         "docx"),
        ("qdrant-client",       "qdrant_client"),
        ("sentence-transformers","sentence_transformers"),
        ("tqdm",                "tqdm"),
    ]:
        try:
            __import__(import_name)
        except ImportError:
            missing.append(pkg)
    if missing:
        print(f"[ERROR] 缺少依赖，请先执行：\n  pip install {' '.join(missing)}")
        sys.exit(1)

_check_dependencies()

# ── 正式导入 ──────────────────────────────────────────────────────────────────
import docx
from qdrant_client import QdrantClient
from qdrant_client.http.models import (
    Distance, VectorParams,
    PointStruct, Filter, FieldCondition, MatchValue,
)
from sentence_transformers import SentenceTransformer
from tqdm import tqdm

# ═════════════════════════════════════════════════════════════════════════════
# 1. 配置
# ═════════════════════════════════════════════════════════════════════════════
@dataclass
class Config:
    # 文档
    docx_path: str = "document.docx"        # Word 文件路径

    # 分块策略
    chunk_size: int     = 500               # 每块目标字符数
    chunk_overlap: int  = 100               # 相邻块重叠字符数
    min_chunk_size: int = 50                # 丢弃过短的碎片

    # 嵌入模型（首次运行自动下载，~90 MB）
    embed_model: str = "sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"

    # ── Qdrant 连接 ──────────────────────────────────────────────────────────
    # 三种模式，按需选一：
    #
    # ① 本地内存（默认，无需任何服务，数据不持久化）
    #      qdrant_host = ":memory:"
    #
    # ② 本地 / 自托管（Docker 等）
    #      qdrant_host = "localhost"  qdrant_port = 6333  qdrant_api_key = None
    #
    # ③ Qdrant Cloud
    #      qdrant_host = "https://xxxx.us-east4-0.gcp.cloud.qdrant.io"
    #      qdrant_port = 6333  (或 443，Cloud 默认 HTTPS)
    #      qdrant_api_key = "your-api-key-here"
    #
    qdrant_host: str        = ":memory:"
    qdrant_port: int        = 6333
    qdrant_api_key: str | None = None       # 自托管无认证时填 None
    qdrant_https: bool      = False         # Cloud / TLS 时设为 True

    collection: str = "docx_rag"

    # 批量写入大小
    batch_size: int = 64

# ═════════════════════════════════════════════════════════════════════════════
# 2. Word 文档解析
# ═════════════════════════════════════════════════════════════════════════════
@dataclass
class DocSection:
    """文档中的一个逻辑段落，含元数据"""
    text: str
    heading: str   = ""        # 所属标题
    para_index: int = 0        # 段落序号
    source_file: str = ""


def parse_docx(path: str) -> list[DocSection]:
    """
    按段落提取文本，保留标题上下文。
    忽略空段落和纯空白段落。
    """
    doc = docx.Document(path)
    sections: list[DocSection] = []
    current_heading = ""

    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        style = para.style.name.lower()
        if "heading" in style:
            current_heading = text
            # 标题本身也作为一个 section 保留（可选）
            sections.append(DocSection(
                text=text,
                heading=text,
                para_index=idx,
                source_file=Path(path).name,
            ))
        else:
            sections.append(DocSection(
                text=text,
                heading=current_heading,
                para_index=idx,
                source_file=Path(path).name,
            ))

    print(f"[解析] 共提取 {len(sections)} 个段落，来自：{path}")
    return sections


# ═════════════════════════════════════════════════════════════════════════════
# 3. 文本分块（滑动窗口 + 句子边界）
# ═════════════════════════════════════════════════════════════════════════════
@dataclass
class Chunk:
    """最终写入向量库的最小单元"""
    chunk_id: str
    text: str
    heading: str
    para_index: int
    chunk_index: int
    source_file: str
    char_count: int = field(init=False)

    def __post_init__(self):
        self.char_count = len(self.text)


def _split_into_sentences(text: str) -> list[str]:
    """按中英文标点切句"""
    pattern = r'(?<=[。！？\.!\?])\s*'
    parts = re.split(pattern, text)
    return [p.strip() for p in parts if p.strip()]


def _make_chunks(text: str, chunk_size: int, overlap: int) -> Generator[str, None, None]:
    """滑动窗口切块，尽量在句子边界处截断"""
    sentences = _split_into_sentences(text)
    buffer = ""
    for sent in sentences:
        if len(buffer) + len(sent) <= chunk_size:
            buffer += sent
        else:
            if buffer:
                yield buffer
            # 用重叠内容作为下一块的开头
            overlap_text = buffer[-overlap:] if overlap > 0 else ""
            buffer = overlap_text + sent
    if buffer:
        yield buffer


def build_chunks(sections: list[DocSection], cfg: Config) -> list[Chunk]:
    """把 DocSection 列表转成 Chunk 列表"""
    all_chunks: list[Chunk] = []

    for sec in sections:
        chunk_texts = list(_make_chunks(sec.text, cfg.chunk_size, cfg.chunk_overlap))
        # 段落本身很短时直接作为一个 chunk
        if not chunk_texts:
            chunk_texts = [sec.text]

        for ci, ct in enumerate(chunk_texts):
            if len(ct) < cfg.min_chunk_size:
                continue
            cid = hashlib.md5(
                f"{sec.source_file}:{sec.para_index}:{ci}:{ct[:20]}".encode()
            ).hexdigest()
            all_chunks.append(Chunk(
                chunk_id=cid,
                text=ct,
                heading=sec.heading,
                para_index=sec.para_index,
                chunk_index=ci,
                source_file=sec.source_file,
            ))

    print(f"[分块] 共生成 {len(all_chunks)} 个 chunk"
          f"（size={cfg.chunk_size}, overlap={cfg.chunk_overlap}）")
    return all_chunks


# ═════════════════════════════════════════════════════════════════════════════
# 4. 向量嵌入
# ═════════════════════════════════════════════════════════════════════════════
def load_embedder(model_name: str) -> SentenceTransformer:
    print(f"[嵌入] 加载模型：{model_name}")
    model = SentenceTransformer(model_name)
    print(f"[嵌入] 向量维度：{model.get_sentence_embedding_dimension()}")
    return model


def embed_chunks(
    chunks: list[Chunk],
    model: SentenceTransformer,
    batch_size: int = 64,
) -> list[list[float]]:
    texts = [c.text for c in chunks]
    print(f"[嵌入] 开始编码 {len(texts)} 个 chunk ...")
    embeddings = model.encode(
        texts,
        batch_size=batch_size,
        show_progress_bar=True,
        normalize_embeddings=True,   # 余弦相似度时建议归一化
    )
    return embeddings.tolist()


# ═════════════════════════════════════════════════════════════════════════════
# 5. Qdrant 操作
# ═════════════════════════════════════════════════════════════════════════════
def get_qdrant_client(cfg: Config) -> QdrantClient:
    if cfg.qdrant_host == ":memory:":
        client = QdrantClient(":memory:")
        print("[Qdrant] 使用内存模式（数据不持久化）")
    else:
        kwargs = dict(
            host=cfg.qdrant_host,
            port=cfg.qdrant_port,
            https=cfg.qdrant_https,
        )
        if cfg.qdrant_api_key:
            kwargs["api_key"] = cfg.qdrant_api_key

        client = QdrantClient(**kwargs)
        auth_hint = "✓ API Key 已配置" if cfg.qdrant_api_key else "无 API Key（本地模式）"
        proto = "https" if cfg.qdrant_https else "http"
        print(f"[Qdrant] 连接至 {proto}://{cfg.qdrant_host}:{cfg.qdrant_port}  {auth_hint}")
    return client


def ensure_collection(client: QdrantClient, cfg: Config, vector_dim: int):
    """若 collection 不存在则创建"""
    existing = [c.name for c in client.get_collections().collections]
    if cfg.collection not in existing:
        client.create_collection(
            collection_name=cfg.collection,
            vectors_config=VectorParams(
                size=vector_dim,
                distance=Distance.COSINE,
            ),
        )
        print(f"[Qdrant] 创建 collection：{cfg.collection}（dim={vector_dim}）")
    else:
        print(f"[Qdrant] collection 已存在：{cfg.collection}")


def upsert_chunks(
    client: QdrantClient,
    cfg: Config,
    chunks: list[Chunk],
    embeddings: list[list[float]],
):
    """批量写入 Qdrant"""
    total = len(chunks)
    print(f"[Qdrant] 开始写入 {total} 个向量，batch_size={cfg.batch_size} ...")

    for start in tqdm(range(0, total, cfg.batch_size), desc="写入进度"):
        end = min(start + cfg.batch_size, total)
        batch_chunks = chunks[start:end]
        batch_vecs   = embeddings[start:end]

        points = [
            PointStruct(
                id=str(uuid.UUID(c.chunk_id)),   # Qdrant 要求 UUID 格式
                vector=vec,
                payload={
                    "text":        c.text,
                    "heading":     c.heading,
                    "para_index":  c.para_index,
                    "chunk_index": c.chunk_index,
                    "source_file": c.source_file,
                    "char_count":  c.char_count,
                },
            )
            for c, vec in zip(batch_chunks, batch_vecs)
        ]
        client.upsert(collection_name=cfg.collection, points=points)

    info = client.get_collection(cfg.collection)
    print(f"[Qdrant] 写入完成，collection 当前向量数：{info.points_count}")


# ═════════════════════════════════════════════════════════════════════════════
# 6. RAG 检索演示
# ═════════════════════════════════════════════════════════════════════════════
def search(
    query: str,
    client: QdrantClient,
    model: SentenceTransformer,
    cfg: Config,
    top_k: int = 5,
    source_file: str | None = None,
) -> list[dict]:
    """
    向量相似度检索，可选按 source_file 过滤。
    返回 top_k 个最相关 chunk 及其 payload。
    """
    query_vec = model.encode(query, normalize_embeddings=True).tolist()

    qdrant_filter = None
    if source_file:
        qdrant_filter = Filter(
            must=[FieldCondition(key="source_file",
                                 match=MatchValue(value=source_file))]
        )

    results = client.search(
        collection_name=cfg.collection,
        query_vector=query_vec,
        limit=top_k,
        query_filter=qdrant_filter,
        with_payload=True,
    )

    hits = []
    for r in results:
        hits.append({
            "score":   round(r.score, 4),
            "heading": r.payload.get("heading", ""),
            "text":    r.payload.get("text", ""),
            "source":  r.payload.get("source_file", ""),
        })
    return hits


# ═════════════════════════════════════════════════════════════════════════════
# 7. 主流程
# ═════════════════════════════════════════════════════════════════════════════
def run_pipeline(cfg: Config):
    print("=" * 60)
    print("  Word → Chunk → Embed → Qdrant  RAG 管道")
    print("=" * 60)

    # ① 检查文件
    if not Path(cfg.docx_path).exists():
        print(f"[WARN] 找不到文件 '{cfg.docx_path}'，将创建演示文档 demo.docx ...")
        _create_demo_docx("demo.docx")
        cfg.docx_path = "demo.docx"

    # ② 解析 Word
    sections = parse_docx(cfg.docx_path)
    if not sections:
        print("[ERROR] 文档为空，退出。")
        return

    # ③ 分块
    chunks = build_chunks(sections, cfg)
    if not chunks:
        print("[ERROR] 分块结果为空，请检查文档内容。")
        return

    # ④ 嵌入
    model = load_embedder(cfg.embed_model)
    embeddings = embed_chunks(chunks, model, cfg.batch_size)

    # ⑤ 写入 Qdrant
    client = get_qdrant_client(cfg)
    ensure_collection(client, cfg, vector_dim=len(embeddings[0]))
    upsert_chunks(client, cfg, chunks, embeddings)

    # ⑥ 演示检索
    print("\n" + "─" * 60)
    demo_queries = [
        "主要内容是什么",
        "核心结论",
    ]
    for q in demo_queries:
        print(f"\n🔍 查询：「{q}」")
        hits = search(q, client, model, cfg, top_k=3)
        for i, h in enumerate(hits, 1):
            print(f"  [{i}] score={h['score']}  heading=《{h['heading']}》")
            print(f"      {h['text'][:120].replace(chr(10), ' ')} ...")

    print("\n✅ 全部流程完成！")
    return client, model, cfg   # 返回供交互式使用


# ═════════════════════════════════════════════════════════════════════════════
# 工具：生成演示 docx（无真实文档时使用）
# ═════════════════════════════════════════════════════════════════════════════
def _create_demo_docx(path: str):
    from docx import Document
    from docx.shared import Pt

    document = Document()
    document.add_heading("RAG 技术概述", level=1)
    document.add_paragraph(
        "检索增强生成（RAG）是一种将信息检索与大语言模型结合的技术框架。"
        "它通过在生成回答之前先从知识库中检索相关文档，显著提升了模型的事实准确性和可追溯性。"
    )
    document.add_heading("核心组件", level=2)
    document.add_paragraph(
        "RAG 系统通常由三个核心模块构成：文档处理模块、向量检索模块和语言生成模块。"
        "文档处理模块负责将原始文本切分为适合嵌入的小块（chunk）。"
    )
    document.add_paragraph(
        "向量检索模块使用嵌入模型将文本和查询转化为高维向量，"
        "并通过余弦相似度等度量在向量数据库中快速找到最相关的片段。"
    )
    document.add_heading("Qdrant 向量数据库", level=2)
    document.add_paragraph(
        "Qdrant 是一款高性能的开源向量数据库，支持过滤、分组和混合检索。"
        "其 Python 客户端易于集成，提供内存模式方便本地调试，"
        "也可通过 Docker 部署为持久化服务。"
    )
    document.add_heading("分块策略", level=2)
    document.add_paragraph(
        "分块大小（chunk size）和重叠量（overlap）是影响 RAG 检索质量的关键参数。"
        "过小的 chunk 会丢失上下文，过大则会引入噪声。"
        "常见策略包括固定长度、句子边界和语义聚类等方法。"
    )
    document.add_heading("结论", level=1)
    document.add_paragraph(
        "RAG 是目前企业知识问答场景中最实用的 LLM 落地方案之一。"
        "结合 Qdrant 等高效向量数据库，可以实现毫秒级的语义检索，"
        "大幅降低幻觉风险并提升答案的可解释性。"
    )
    document.save(path)
    print(f"[演示] 已生成演示文档：{path}")


# ═════════════════════════════════════════════════════════════════════════════
# 入口
# ═════════════════════════════════════════════════════════════════════════════
def parse_args() -> Config:
    parser = argparse.ArgumentParser(description="Word → Qdrant RAG 管道")
    parser.add_argument("--file",        default="/Users/longxuebin/Downloads/我的下载toby/20260302B端假期内测学校2025级高一年级错题试卷集学生卷 (1)/语文错题集学生卷.docx", help="Word 文件路径")
    parser.add_argument("--collection",  default="docx_rag",      help="Qdrant collection 名称")
    parser.add_argument("--chunk-size",  type=int, default=500,   help="每块目标字符数")
    parser.add_argument("--overlap",     type=int, default=100,   help="相邻块重叠字符数")
    parser.add_argument("--qdrant-host", default="https://95abf2e0-7725-4e71-a1c5-5c7deae88269.us-west-1-0.aws.cloud.qdrant.io:6333",      help="Qdrant 主机（':memory:' = 内存模式）")
    parser.add_argument("--qdrant-port", type=int, default=6333,  help="Qdrant 端口")
    parser.add_argument("--api-key",     default="eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJhY2Nlc3MiOiJtIiwic3ViamVjdCI6ImFwaS1rZXk6MzQwYjZlNTUtNmQwYi00MzFjLWIzODQtYTU5YzE5NDc1ZDg0In0.evMEbWRPlSH8pM83HG88M0JGBoB3Ih_0nTGh3ZQSd80",            help="Qdrant API Key（Cloud 或开启认证的自托管）")
    parser.add_argument("--https",       action="store_true",     help="使用 HTTPS 连接（Qdrant Cloud 必须加此项）")
    parser.add_argument("--model",
        default="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2",
        help="HuggingFace 嵌入模型名称")
    args = parser.parse_args()

    return Config(
        docx_path=args.file,
        collection=args.collection,
        chunk_size=args.chunk_size,
        chunk_overlap=args.overlap,
        qdrant_host=args.qdrant_host,
        qdrant_port=args.qdrant_port,
        qdrant_api_key=args.api_key,
        qdrant_https=args.https,
        embed_model=args.model,
    )


if __name__ == "__main__":
    cfg = parse_args()
    run_pipeline(cfg)