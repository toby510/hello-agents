from qdrant_client import QdrantClient
from qdrant_client.models import (
    Distance,
    VectorParams,
    PointStruct,
    Filter,
    FieldCondition,
    MatchValue,
    PayloadSchemaType,
)

COLLECTION_NAME = "demo_collection"
VECTOR_SIZE = 4

qdrant_client = QdrantClient(
    url="https://95abf2e0-7725-4e71-a1c5-5c7deae88269.us-west-1-0.aws.cloud.qdrant.io:6333",
    api_key="eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJhY2Nlc3MiOiJtIiwic3ViamVjdCI6ImFwaS1rZXk6MzQwYjZlNTUtNmQwYi00MzFjLWIzODQtYTU5YzE5NDc1ZDg0In0.evMEbWRPlSH8pM83HG88M0JGBoB3Ih_0nTGh3ZQSd80"
)


def show_collections():
    print("=== 1. 查看所有 Collections ===")
    collections = qdrant_client.get_collections()
    print(collections)
    print()


def create_collection():
    print("=== 2. 创建 Collection ===")
    # 如果已存在则先删除，保证可重复运行
    if qdrant_client.collection_exists(COLLECTION_NAME):
        qdrant_client.delete_collection(COLLECTION_NAME)
        print(f"旧 Collection '{COLLECTION_NAME}' 已删除")

    qdrant_client.create_collection(
        collection_name=COLLECTION_NAME,
        vectors_config=VectorParams(size=VECTOR_SIZE, distance=Distance.COSINE),
    )
    # 为过滤字段创建索引，否则过滤搜索会报错
    qdrant_client.create_payload_index(
        collection_name=COLLECTION_NAME,
        field_name="category",
        field_type=PayloadSchemaType.KEYWORD,
    )
    print(f"Collection '{COLLECTION_NAME}' 创建成功，并已为 'category' 字段创建索引")
    print()


def insert_points():
    print("=== 3. 插入 Points (Create) ===")
    points = [
        PointStruct(
            id=1,
            vector=[0.1, 0.2, 0.3, 0.4],
            payload={"name": "apple", "category": "fruit", "price": 5.5},
        ),
        PointStruct(
            id=2,
            vector=[0.3, 0.3, 0.3, 0.3],
            payload={"name": "banana2", "category": "fruit", "price": 3.0},
        ),
        PointStruct(
            id=3,
            vector=[0.8, 0.9, 0.1, 0.2],
            payload={"name": "carrot", "category": "vegetable", "price": 2.5},
        ),
        PointStruct(
            id=4,
            vector=[0.85, 0.88, 0.12, 0.18],
            payload={"name": "broccoli", "category": "vegetable", "price": 4.0},
        ),
    ]
    operation_info = qdrant_client.upsert(
        collection_name=COLLECTION_NAME,
        wait=True,
        points=points,
    )
    print(f"插入结果: {operation_info.status}")
    print()


def read_points():
    print("=== 4. 读取 Points (Read) ===")
    # 按 ID 查询
    points = qdrant_client.retrieve(
        collection_name=COLLECTION_NAME,
        ids=[1, 2],
        with_payload=True,
        with_vectors=True,
    )
    for p in points:
        print(f"  ID={p.id}, payload={p.payload}, vector={p.vector}")
    print()


def update_point():
    print("=== 5. 更新 Point (Update) ===")
    # 更新 payload
    qdrant_client.set_payload(
        collection_name=COLLECTION_NAME,
        payload={"price": 6.0, "on_sale": True},
        points=[1],
    )
    # 也可以覆盖整个 point（upsert）
    qdrant_client.upsert(
        collection_name=COLLECTION_NAME,
        points=[
            PointStruct(
                id=2,
                vector=[0.25, 0.15, 0.35, 0.45],
                payload={"name": "banana3", "category": "fruit", "price": 3.5},
            )
        ],
    )
    print("Point 1 payload 更新完成, Point 2 vector 和 payload 更新完成")

    # 验证更新结果
    points = qdrant_client.retrieve(
        collection_name=COLLECTION_NAME, ids=[1, 2], with_payload=True
    )
    for p in points:
        print(f"  ID={p.id}, payload={p.payload}")
    print()


def search_similar():
    print("=== 6. 向量搜索 (Search) ===")
    query_vector = [0.82, 0.87, 0.11, 0.19]  # 接近 carrot / broccoli
    response = qdrant_client.query_points(
        collection_name=COLLECTION_NAME,
        query=query_vector,
        limit=3,
        with_payload=True,
    )
    for r in response.points:
        print(f"  ID={r.id}, score={r.score:.4f}, payload={r.payload}")
    print()


def search_with_filter():
    print("=== 7. 带过滤条件的搜索 (Search + Filter) ===")
    query_vector = [0.15, 0.15, 0.35, 0.35]
    response = qdrant_client.query_points(
        collection_name=COLLECTION_NAME,
        query=query_vector,
        query_filter=Filter(
            must=[
                FieldCondition(
                    key="category",
                    match=MatchValue(value="fruit"),
                )
            ]
        ),
        limit=3,
        with_payload=True,
    )
    for r in response.points:
        print(f"  ID={r.id}, score={r.score:.4f}, payload={r.payload}")
    print()


def count_points():
    print("=== 8. 统计 Points ===")
    count = qdrant_client.count(collection_name=COLLECTION_NAME)
    print(f"总点数: {count.count}")
    print()


def delete_points():
    print("=== 9. 删除 Points (Delete) ===")
    qdrant_client.delete(
        collection_name=COLLECTION_NAME,
        points_selector=[1],
    )
    print("已删除 ID=1 的 point")

    count = qdrant_client.count(collection_name=COLLECTION_NAME)
    print(f"删除后总点数: {count.count}")
    print()


def scroll_points():
    print("=== 10. 滚动遍历 (Scroll) ===")
    records, next_page_offset = qdrant_client.scroll(
        collection_name=COLLECTION_NAME,
        limit=10,
        with_payload=True,
    )
    for r in records:
        print(f"  ID={r.id}, payload={r.payload}")
    print()


def cleanup():
    print("=== 11. 清理 Collection ===")
    qdrant_client.delete_collection(COLLECTION_NAME)
    print(f"Collection '{COLLECTION_NAME}' 已删除")
    print()


# ===================== 文档向量化 RAG 扩展 =====================

DOC_COLLECTION = "doc_chunks"
DOC_VECTOR_SIZE = 384  # all-MiniLM-L6-v2


def _load_text(file_path: str) -> str:
    """根据后缀读取 .docx 或 .pdf 的纯文本。"""
    ext = file_path.lower().split(".")[-1]
    if ext == "docx":
        from docx import Document

        doc = Document(file_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        return "\n".join(paragraphs)
    elif ext == "pdf":
        import pymupdf

        parts = []
        with pymupdf.open(file_path) as doc:
            for page in doc:
                parts.append(page.get_text())
        return "\n".join(parts)
    else:
        raise ValueError(f"不支持的文件格式: {ext}，仅支持 .docx 和 .pdf")


def _preprocess(text: str) -> str:
    """文本预处理：去空行、去多余空格、清理格式符。"""
    import re

    text = re.sub(r"\s+", " ", text)
    return text.strip()


def _chunk_text(text: str, chunk_size: int = 500, chunk_overlap: int = 50) -> list:
    """按固定长度分块，允许重叠。"""
    chunks = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        chunks.append(text[start:end])
        if end == len(text):
            break
        start += chunk_size - chunk_overlap
    return chunks


def _get_embedding_model():
    """懒加载 embedding 模型（避免影响原有 demo 启动速度）。"""
    import os

    os.environ.setdefault("HF_ENDPOINT", "https://hf-mirror.com")
    from sentence_transformers import SentenceTransformer

    if not hasattr(_get_embedding_model, "_model"):
        _get_embedding_model._model = SentenceTransformer("all-MiniLM-L6-v2")
    return _get_embedding_model._model


def _ensure_doc_collection():
    """确保文档向量 collection 及 payload 索引存在。"""
    if qdrant_client.collection_exists(DOC_COLLECTION):
        return
    qdrant_client.create_collection(
        collection_name=DOC_COLLECTION,
        vectors_config=VectorParams(size=DOC_VECTOR_SIZE, distance=Distance.COSINE),
    )
    qdrant_client.create_payload_index(
        collection_name=DOC_COLLECTION,
        field_name="source_file",
        field_type=PayloadSchemaType.KEYWORD,
    )
    qdrant_client.create_payload_index(
        collection_name=DOC_COLLECTION,
        field_name="chunk_index",
        field_type=PayloadSchemaType.INTEGER,
    )
    print(f"Collection '{DOC_COLLECTION}' 创建成功")


def ingest_document(file_path: str, chunk_size: int = 500, chunk_overlap: int = 50):
    """
    1. 加载文档（.docx / .pdf）
    2. 文本预处理
    3. 文本分块（Chunk）
    4. 嵌入向量化
    5. 存入向量库
    """
    import os

    print(f"=== 开始处理文档: {file_path} ===")

    raw_text = _load_text(file_path)
    print(f"原始文本长度: {len(raw_text)} 字符")

    clean_text = _preprocess(raw_text)
    print(f"预处理后长度: {len(clean_text)} 字符")

    chunks = _chunk_text(clean_text, chunk_size, chunk_overlap)
    print(f"分块数量: {len(chunks)} (size={chunk_size}, overlap={chunk_overlap})")

    model = _get_embedding_model()
    embeddings = model.encode(chunks, show_progress_bar=True)
    print(f"向量化完成，维度: {embeddings.shape[1]}")

    _ensure_doc_collection()

    import uuid

    base_name = os.path.basename(file_path)
    points = []
    for idx, (chunk, vec) in enumerate(zip(chunks, embeddings)):
        point_id = str(uuid.uuid5(uuid.NAMESPACE_DNS, f"{base_name}_{idx}"))
        points.append(
            PointStruct(
                id=point_id,
                vector=vec.tolist(),
                payload={
                    "text": chunk,
                    "source_file": file_path,
                    "chunk_index": idx,
                    "total_chunks": len(chunks),
                },
            )
        )

    qdrant_client.upsert(collection_name=DOC_COLLECTION, wait=True, points=points)
    print(f"成功写入 {len(points)} 个 chunk 到 '{DOC_COLLECTION}'\n")


def search_document_chunks(query_text: str, limit: int = 10):
    """
    将查询文本向量化后，在 doc_chunks 中检索相似度前 N 的 chunk。
    """
    print(f"=== 查询: '{query_text}' ===")

    model = _get_embedding_model()
    query_vector = model.encode([query_text])[0].tolist()

    response = qdrant_client.query_points(
        collection_name=DOC_COLLECTION,
        query=query_vector,
        limit=limit,
        with_payload=True,
    )

    print(f"返回结果数: {len(response.points)}")
    for r in response.points:
        payload = r.payload
        snippet = payload.get("text", "")[:120].replace("\n", " ")
        print(f"  score={r.score:.4f} | chunk_index={payload.get('chunk_index')}/{payload.get('total_chunks')} | text={snippet}...")
    print()
    return response.points


def doc_test():
    """文档向量化全流程测试（生成临时文件 -> 入库 -> 检索）。"""
    import os
    import pymupdf
    from docx import Document

    test_docx = "/tmp/test_demo.docx"
    test_pdf = "/tmp/test_demo.pdf"

    # 生成测试 .docx
    doc = Document()
    doc.add_paragraph(
        "人工智能（Artificial Intelligence，AI）是指由人制造出来的系统所表现出来的智能。\n"
        "通常人工智能是指通过普通计算机程序来呈现人类智能的技术。\n"
        "该词也指出研究这样的智能系统是否能够实现，以及如何实现。\n"
        "人工智能于一般教材中的定义领域是「智能主体（intelligent agent）的研究与设计」。\n"
        "智能主体指一个可以观察周遭环境并作出行动以达致目标的系统。\n"
        "约翰·麦卡锡于1955年的定义是「制造智能机器，特别是智能计算机程序的科学和工程」。\n"
        "AI的核心问题包括建构能够跟人类似甚至超卓的推理、知识、规划、学习、交流、感知、移物、使用工具和操控机械的能力等。\n"
        "人工智能已经在多个领域中得到了广泛应用，包括医疗诊断、金融分析、自动驾驶、自然语言处理等。\n"
        "机器学习是人工智能的一个重要分支，它使计算机能够在没有明确编程的情况下从数据中学习。\n"
        "深度学习是机器学习的一个子集，使用多层神经网络来模拟人脑的工作方式。\n"
        "神经网络由相互连接的节点层组成，每个连接都有一个相关的权重和阈值。\n"
        "卷积神经网络（CNN）特别适用于图像识别和处理任务。\n"
        "循环神经网络（RNN）则擅长处理序列数据，如时间序列或自然语言。\n"
        "Transformer 架构是目前自然语言处理领域的主流，它使用自注意力机制来捕捉序列中的长距离依赖关系。\n"
        "大型语言模型（LLM）如 GPT、Claude 等基于 Transformer 架构，能够生成连贯且上下文相关的文本。\n"
        "向量数据库在 RAG（检索增强生成）架构中扮演着关键角色，用于存储和检索文本的向量表示。\n"
        "Qdrant 是一个高性能的向量数据库，支持高效的相似度搜索和过滤。"
    )
    doc.save(test_docx)
    print(f"已生成测试文件: {test_docx}")

    # 生成测试 .pdf
    pdf_doc = pymupdf.open()
    page = pdf_doc.new_page()
    content = (
        "向量数据库是一种专门设计用于存储和查询高维向量的数据库系统。\n"
        "与传统关系型数据库不同，向量数据库优化了相似度搜索操作。\n"
        "常见的相似度度量包括欧氏距离、余弦相似度和点积。\n"
        "近似最近邻（ANN）搜索算法如 HNSW、IVF 等被广泛用于加速大规模向量检索。\n"
        "在推荐系统中，向量数据库可以存储用户和物品的嵌入向量，实现实时个性化推荐。\n"
        "图像检索系统使用向量数据库存储图像特征向量，支持以图搜图功能。\n"
        "文本语义搜索将查询和文档编码为向量，通过向量相似度找到语义相关的文档。\n"
        "多模态搜索允许使用一种模态（如文本）查询另一种模态（如图像）的数据。\n"
        "Qdrant 使用 Rust 编写，提供了高性能的向量索引和查询能力。\n"
        "它还支持 payload 过滤、分布式部署和混合搜索等高级功能。"
    )
    page.insert_text((72, 72), content, fontsize=12)
    pdf_doc.save(test_pdf)
    pdf_doc.close()
    print(f"已生成测试文件: {test_pdf}")

    # 入库
    ingest_document(test_docx, chunk_size=300, chunk_overlap=30)
    ingest_document(test_pdf, chunk_size=300, chunk_overlap=30)

    # 检索测试
    search_document_chunks("人工智能是什么", limit=10)
    search_document_chunks("向量数据库的作用", limit=10)

    # 清理临时文件
    os.remove(test_docx)
    os.remove(test_pdf)
    print("测试完成，临时文件已清理")


def main():
    show_collections()
    create_collection()
    insert_points()
    read_points()
    update_point()
    search_similar()
    search_with_filter()
    count_points()
    delete_points()
    scroll_points()
    # cleanup()
    print("所有操作执行完毕！")


if __name__ == "__main__":
    # 原有 demo
    # main()

    # 文档向量化测试
    doc_test()
